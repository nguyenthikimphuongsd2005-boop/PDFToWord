import os
import shutil
import fitz
from pdf2docx import Converter
from PyPDF2 import PdfReader
from pdf2image import convert_from_path
import pytesseract
from docx import Document
import cv2
import numpy as np
from PIL import Image

def preprocess_image(pil_image):
    """
    Làm sạch ảnh trước khi OCR để tăng độ chính xác:
    - Chuyển về đen trắng (Grayscale)
    - Tăng độ tương phản (Thresholding)
    - Khử nhiễu
    """
    # Chuyển PIL Image -> OpenCV Image (numpy array)
    img = np.array(pil_image) 
    
    # 1. Chuyển sang Gray
    gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
    
    # 2. Khử nhiễu (Denoise)
    gray = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)

    # 3. Nhị phân hóa (Thresholding - Otsu's method)
    # Giúp chữ tách biệt hẳn khỏi nền
    _, thresh = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

    # Chuyển ngược lại về PIL Image để đưa vào Tesseract
    return Image.fromarray(thresh)

# CẤU HÌNH TESSERACT (Cập nhật đường dẫn server của bạn tại đây)
# Nếu chạy trên Docker hoặc Linux server thì thường không cần dòng này nếu đã cài vào PATH
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


class PDFProcessor:
    @staticmethod
    def parse_page_range(page_str, max_pages):
        if not page_str or not page_str.strip():
            return None
        pages = set()
        parts = page_str.split(",")
        for part in parts:
            part = part.strip()
            if "-" in part:
                try:
                    start, end = map(int, part.split("-"))
                    start = max(1, start)
                    end = min(max_pages, end)
                    if start <= end:
                        for p in range(start - 1, end):
                            pages.add(p)
                except ValueError:
                    continue
            else:
                try:
                    p = int(part)
                    if 1 <= p <= max_pages:
                        pages.add(p - 1)
                except ValueError:
                    continue
        return sorted(list(pages))

    @staticmethod
    def is_scanned_pdf(pdf_path):
        try:
            reader = PdfReader(pdf_path)
            text = ""
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted
            return len(text.strip()) < 50
        except:
            return True

    @staticmethod    
    def ocr_convert(pdf_path, docx_path, page_indices=None):
        try:
            images = convert_from_path(pdf_path)  # Cần cài đặt Poppler
            doc = Document()

            selected_images = (
                [images[i] for i in page_indices if i < len(images)]
                if page_indices
                else images
            )

            for i, img in enumerate(selected_images):
                text = pytesseract.image_to_string(img, lang="vie+eng")
                doc.add_heading(f"Trang {i+1} (OCR)", level=2)
                doc.add_paragraph(text)
                doc.add_page_break()

            doc.save(docx_path)
            return True, "OCR Thành công"
        except Exception as e:
            print(f"Lỗi OCR: {e}")
            return False, str(e)

    @staticmethod
    def convert_to_word(pdf_file, output_path, page_range_str=None, use_ocr=False):
        try:
            reader = PdfReader(pdf_file)
            total_pages = len(reader.pages)
            page_indices = PDFProcessor.parse_page_range(page_range_str, total_pages)

            should_run_ocr = False
            if use_ocr:
                if PDFProcessor.is_scanned_pdf(pdf_file):
                    should_run_ocr = True

            if should_run_ocr:
                return PDFProcessor.ocr_convert(pdf_file, output_path, page_indices)

            # Chế độ thường
            cv = Converter(pdf_file)
            if page_indices:
                cv.convert(output_path, pages=page_indices)
            else:
                cv.convert(output_path)
            cv.close()
            return True, "Thành công"
        except Exception as e:
            return False, str(e)

    @staticmethod
    def convert_to_text(pdf_file, output_path, page_range_str=None):
        try:
            reader = PdfReader(pdf_file)
            total_pages = len(reader.pages)
            page_indices = PDFProcessor.parse_page_range(page_range_str, total_pages)

            text_content = ""
            pages_to_process = (
                page_indices if page_indices is not None else range(total_pages)
            )

            for i in pages_to_process:
                if i < total_pages:
                    extracted = reader.pages[i].extract_text()
                    if extracted:
                        text_content += f"--- Trang {i+1} ---\n{extracted}\n\n"

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(text_content)
            return True, "Thành công"
        except Exception as e:
            return False, str(e)
